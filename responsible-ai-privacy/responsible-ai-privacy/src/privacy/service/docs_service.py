# MIT license https://opensource.org/licenses/MIT
# Copyright 2024 Infosys Ltd
# 
# Permission is hereby granted, free of charge, to any person obtaining a copy of this software and associated documentation files (the "Software"), to deal in the Software without restriction, including without limitation the rights to use, copy, modify, merge, publish, distribute, sublicense, and/or sell copies of the Software, and to permit persons to whom the Software is furnished to do so, subject to the following conditions:
# 
# The above copyright notice and this permission notice shall be included in all copies or substantial portions of the Software.
# 
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE SOFTWARE.



import base64
import io
from typing import Tuple
from PIL import Image
from privacy.service.imagePrivacy import AttributeDict, ImagePrivacy
from privacy.service.api_req import ApiCall
from privacy.service.textPrivacy import TextPrivacy
from privacy.service.__init__ import *
import numpy as np
from privacy.config.logger import request_id_var
from privacy.config.logger import CustomLogger
import os
import threading
import uuid
from docx import Document
from docx.shared import Inches
from unidecode import unidecode
import tempfile

log = CustomLogger()
error_dict = {}

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class DOCService:
    def processImages(paragraph, run, payload, uid):
        try:
            request_id_var.set(uid)
            blip = run._element.xpath('.//pic:blipFill/a:blip')[0]
            rId = blip.get(qn('r:embed'))
            log.debug(f"Relationship ID: {rId}")
            
            # Get the image part using the relationship ID
            document_part = run.part
            image_part = document_part.related_parts[rId]
            img_bytes = image_part.blob
            log.debug(f"Image bytes length: {len(img_bytes)}")
            
            if len(img_bytes) < 700:
                log.debug("Image is too small to process.")
                return None
            
            imgd = io.BytesIO(img_bytes)
            payload["image"] = AttributeDict({"file": imgd})
            payload["piiEntitiesToBeRedacted"] = None
            resImage = ImagePrivacy.image_anonymize(AttributeDict(payload))
            log.debug(f"Anonymized image response: {resImage}")
            resImg = base64.b64decode(resImage)
            img_stream = io.BytesIO(resImg)
            img = Image.open(img_stream)
            
            # Use tempfile to create a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_file:
                img_path = tmp_file.name
                img.save(img_path)
            
            run.clear()
            run.add_picture(img_path, width=Inches(1.25))
            os.remove(img_path)
        except Exception as e:
            log.error(str(e))
            log.error("Line No:" + str(e.__traceback__.tb_lineno))
            log.error(str(e.__traceback__.tb_frame))
            if request_id_var.get() not in error_dict:
                error_dict[request_id_var.get()] = []
            error_dict[request_id_var.get()].append({"UUID": request_id_var.get(), "function": "DOCMASKMainFunction", "msg": str(e.__class__.__name__), "description": str(e) + "Line No:" + str(e.__traceback__.tb_lineno)})
            raise Exception(e)

    def editText(text, i, run):
        request_id_var.set("editText")
        log.debug(str(text[i.start:i.end]) + ":" + str(i.entity_type))
        run.text = run.text.replace(text[i.start:i.end], f"<{i.entity_type}>")

    def processText(paragraph, payload, uid):
        try:
            request_id_var.set(uid)
            for run in paragraph.runs:
                text = unidecode(run.text)
                accDetails = None
                if payload.portfolio is not None:
                    accDetails = AttributeDict({"portfolio": payload.portfolio, "account": payload.account})
                res = TextPrivacy.textAnalyze(text=text, accName=accDetails, exclusion=payload.exclusion.split(',') if payload.exclusion is not None else [])
                res = anonymizer._remove_conflicts_and_get_text_manipulation_data(res, (ConflictResolutionStrategy.MERGE_SIMILAR_OR_CONTAINED))
                res = anonymizer._merge_entities_with_whitespace_between(text, res)
                resThreads = []
                for i in res:
                    thread = threading.Thread(target=DOCService.editText, args=(text, i, run))
                    thread.start()
                    resThreads.append(thread)
                for thread in resThreads:
                    thread.join()
        except Exception as e:
            log.error(str(e))
            log.error("Line No:" + str(e.__traceback__.tb_lineno))
            log.error(str(e.__traceback__.tb_frame))
            if request_id_var.get() not in error_dict:
                error_dict[request_id_var.get()] = []
            error_dict[request_id_var.get()].append({"UUID": request_id_var.get(), "function": "DOCMASKMainFunction", "msg": str(e.__class__.__name__), "description": str(e) + "Line No:" + str(e.__traceback__.tb_lineno)})
            raise Exception(e)

    def mask_doc(payload):
        try:
            log.debug("payload:-" + str(payload))
            id = uuid.uuid4().hex
            request_id_var.set(id)

            if payload.portfolio is not None or payload.account is not None:
                response_value = ApiCall.request(AttributeDict({"portfolio": payload.portfolio, "account": payload.account}))
                if response_value is None:
                    return None

            doc_file = Document(io.BytesIO(payload.file.file.read()))

            for paragraph in doc_file.paragraphs:
                threads = []
                thread = threading.Thread(target=DOCService.processText, args=(paragraph, payload, id))
                thread.start()
                threads.append(thread)
                for run in paragraph.runs:
                    if run._element.xpath('.//pic:blipFill/a:blip'):
                        thread = threading.Thread(target=DOCService.processImages, args=(paragraph, run, payload, id))
                        thread.start()
                        threads.append(thread)
                for thread in threads:
                    thread.join()

            doc_bytes = io.BytesIO()
            doc_file.save(doc_bytes)
            doc_bytes.seek(0)

            return doc_bytes
        except Exception as e:
            log.error(str(e))
            log.error("Line No:" + str(e.__traceback__.tb_lineno))
            log.error(str(e.__traceback__.tb_frame))
            if request_id_var.get() not in error_dict:
                error_dict[request_id_var.get()] = []
            error_dict[request_id_var.get()].append({"UUID": request_id_var.get(), "function": "DOCMASKMainFunction", "msg": str(e.__class__.__name__), "description": str(e) + "Line No:" + str(e.__traceback__.tb_lineno)})
            raise Exception(e)
        
